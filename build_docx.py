from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colours ──────────────────────────────────────────────────────────────
PURPLE      = RGBColor(0x64, 0x3C, 0xFF)   # Anchor primary
PURPLE_DARK = RGBColor(0x3A, 0x1E, 0xCC)
PURPLE_SOFT = RGBColor(0xF0, 0xEB, 0xFF)   # light tint for table headers
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
NEAR_BLACK  = RGBColor(0x12, 0x0D, 0x2A)
GRAY_DARK   = RGBColor(0x44, 0x40, 0x5A)
GRAY_MID    = RGBColor(0x88, 0x84, 0x9E)
GRAY_LIGHT  = RGBColor(0xF4, 0xF2, 0xFA)
RISK_RED    = RGBColor(0xFF, 0x4D, 0x4D)
RISK_AMBER  = RGBColor(0xFF, 0xA5, 0x00)
RISK_GREEN  = RGBColor(0x22, 0xC5, 0x5E)

FONT_NAME = "Calibri"

doc = Document()

# ── Page margins ───────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ────────────────────────────────────────────────────────────────────
def set_font(run, size, bold=False, italic=False, color=None, name=FONT_NAME):
    run.font.name  = name
    run.font.size  = Pt(size)
    run.font.bold  = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color

def para_spacing(para, before=0, after=0, line=None):
    pf = para.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after  = Pt(after)
    if line:
        pf.line_spacing = Pt(line)

def shade_cell(cell, hex_color: str):
    """Fill a table cell with a solid colour (hex without #)."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def cell_margins(cell, top=80, bottom=80, left=120, right=120):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'),    str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_cell_border(cell, **kwargs):
    """kwargs: top, bottom, left, right — each a dict with color, sz, val."""
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            cfg  = kwargs[edge]
            node = OxmlElement(f'w:{edge}')
            node.set(qn('w:val'),   cfg.get('val', 'single'))
            node.set(qn('w:sz'),    str(cfg.get('sz', 4)))
            node.set(qn('w:space'), '0')
            node.set(qn('w:color'), cfg.get('color', 'auto'))
            tcBorders.append(node)
    tcPr.append(tcBorders)

def no_borders(table):
    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell,
                top={'val':'nil'}, bottom={'val':'nil'},
                left={'val':'nil'}, right={'val':'nil'})

def thin_bottom(cell, color='DDDBEA'):
    set_cell_border(cell,
        top={'val':'nil'}, left={'val':'nil'}, right={'val':'nil'},
        bottom={'val':'single', 'sz': 4, 'color': color})

def add_horizontal_rule(doc, color='643CFF', thickness=12):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pb  = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(thickness))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), color)
    pb.append(bot)
    pPr.append(pb)
    para_spacing(p, before=4, after=4)
    return p

def page_break(doc):
    p  = doc.add_paragraph()
    r  = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r._r.append(br)
    para_spacing(p, before=0, after=0)

# ══════════════════════════════════════════════════════════════════════════════
# COVER PAGE
# ══════════════════════════════════════════════════════════════════════════════
def add_cover(doc):
    # Spacer
    for _ in range(4):
        sp = doc.add_paragraph()
        para_spacing(sp, before=0, after=0)

    # Logo-like anchor symbol
    anchor_p = doc.add_paragraph()
    anchor_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = anchor_p.add_run("⚓")
    r.font.size  = Pt(36)
    r.font.color.rgb = PURPLE
    para_spacing(anchor_p, before=0, after=6)

    # Product name
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("Anchor SMS Copilot")
    set_font(r, 34, bold=True, color=NEAR_BLACK)
    para_spacing(title_p, before=0, after=4)

    # Tagline
    tag_p = doc.add_paragraph()
    tag_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = tag_p.add_run("Product Concept & Prototype Documentation")
    set_font(r, 13, color=GRAY_MID)
    para_spacing(tag_p, before=0, after=2)

    # Thin purple rule
    rule = doc.add_paragraph()
    rule.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = rule.add_run("─" * 28)
    r.font.color.rgb = PURPLE
    r.font.size = Pt(11)
    para_spacing(rule, before=6, after=6)

    # Sub-line
    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub_p.add_run("Billing Platform  ·  AI-Powered Interface  ·  Freelancer-First")
    set_font(r, 10, italic=True, color=GRAY_MID)
    para_spacing(sub_p, before=0, after=0)

    page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION HEADING helpers
# ══════════════════════════════════════════════════════════════════════════════
def section_heading(doc, text, level=1, numbered=None):
    """level 1 = big section, level 2 = subsection, level 3 = small heading."""
    p = doc.add_paragraph()
    if level == 1:
        # coloured pill-style heading
        r = p.add_run(f"  {text}  " if not numbered else f"  {numbered}   {text}  ")
        set_font(r, 15, bold=True, color=WHITE)
        # background via highlight isn't easy; use shading on a 1-cell table instead
        p.clear()  # discard — rebuild as table trick
        tbl = doc.add_table(rows=1, cols=1)
        no_borders(tbl)
        cell = tbl.cell(0, 0)
        shade_cell(cell, '643CFF')
        cell_margins(cell, top=100, bottom=100, left=160, right=160)
        cp = cell.paragraphs[0]
        r  = cp.add_run(text.upper())
        set_font(r, 12, bold=True, color=WHITE)
        para_spacing(cp, before=0, after=0)
        # add spacing after table
        sp = doc.add_paragraph()
        para_spacing(sp, before=2, after=2)

    elif level == 2:
        r = p.add_run(text)
        set_font(r, 13, bold=True, color=PURPLE_DARK)
        para_spacing(p, before=14, after=3)
        # underline via border
        pPr = p._p.get_or_add_pPr()
        pb  = OxmlElement('w:pBdr')
        bot = OxmlElement('w:bottom')
        bot.set(qn('w:val'),   'single')
        bot.set(qn('w:sz'),    '4')
        bot.set(qn('w:space'), '1')
        bot.set(qn('w:color'), 'C4BAFF')
        pb.append(bot)
        pPr.append(pb)

    elif level == 3:
        r = p.add_run(text)
        set_font(r, 11, bold=True, color=PURPLE)
        para_spacing(p, before=10, after=2)

def body(doc, text, indent=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, 10.5, color=GRAY_DARK)
    para_spacing(p, before=0, after=5, line=14)
    if indent:
        p.paragraph_format.left_indent = Cm(0.7)
    return p

def bullet(doc, text, sub=False):
    p = doc.add_paragraph(style='List Bullet')
    # strip default bullet style formatting — add our own
    p.clear()
    pf = p.paragraph_format
    pf.left_indent   = Cm(1.2) if not sub else Cm(2.0)
    pf.first_line_indent = Cm(-0.45)
    pf.space_before  = Pt(2)
    pf.space_after   = Pt(2)
    bullet_char = "•" if not sub else "–"
    rb = p.add_run(f"{bullet_char}  ")
    set_font(rb, 10.5, bold=True, color=PURPLE)
    r  = p.add_run(text)
    set_font(r, 10.5, color=GRAY_DARK)
    return p

def bold_inline(para, label, rest, label_color=NEAR_BLACK, rest_color=GRAY_DARK):
    r1 = para.add_run(label)
    set_font(r1, 10.5, bold=True, color=label_color)
    r2 = para.add_run(rest)
    set_font(r2, 10.5, color=rest_color)

def callout_box(doc, title, body_text, bg='F0EBFF', accent='643CFF'):
    tbl = doc.add_table(rows=1, cols=1)
    no_borders(tbl)
    cell = tbl.cell(0, 0)
    shade_cell(cell, bg)
    cell_margins(cell, top=120, bottom=120, left=180, right=180)
    set_cell_border(cell,
        top   ={'val':'single','sz':12,'color':accent},
        bottom={'val':'nil'}, left={'val':'nil'}, right={'val':'nil'})
    cp  = cell.paragraphs[0]
    rt  = cp.add_run(f"{title}  ")
    set_font(rt, 10.5, bold=True, color=RGBColor(
        int(accent[0:2],16), int(accent[2:4],16), int(accent[4:6],16)))
    rb  = cp.add_run(body_text)
    set_font(rb, 10.5, color=GRAY_DARK)
    para_spacing(cp, before=0, after=0)
    sp = doc.add_paragraph()
    para_spacing(sp, before=2, after=2)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 1 — OVERVIEW
# ══════════════════════════════════════════════════════════════════════════════
add_cover(doc)

section_heading(doc, "Overview")
body(doc,
    "Anchor is a billing and payments platform designed to help small businesses and "
    "freelancers manage workflows and get paid faster. However, even with great tools, "
    "solo freelancers and small agencies face significant emotional and administrative "
    "hurdles before a payment is even initiated.")
body(doc,
    "The Anchor SMS Copilot is an AI-powered interface that removes this friction. It "
    "allows users to draft, amend, and send agreements using natural language via standard "
    "text messaging — moving the billing process from a \\"desk chore\\" to a "
    "\\"real-time conversation.\\"")

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 2 — PAIN POINTS
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Pain Points")

pains = [
    ("Getting Paid — The Emotional Tax",
     [
         "The Fear: Late or non-payment is the #1 stressor. Even with contracts, clients delay or \"forget.\"",
         "The Friction: Chasing money is emotionally exhausting. Collection calls and \"just checking in\" "
         "emails feel like begging, damaging the freelancer-client relationship.",
         "Manual Enforcement: Milestone structuring (like the 50/25/25 rule) is often manual and "
         "awkward to bring up in casual negotiations.",
     ]),
    ("Cash Flow Unpredictability",
     [
         "The Anxiety Gap: Freelancers often don't know when money will actually land, making personal "
         "financial planning impossible.",
         "The Initiation Lag: If a project ends on a Friday but the freelancer doesn't invoice until "
         "Monday, they've already lost three days of cash flow.",
     ]),
    ("Agreement → Billing Friction",
     [
         "The Translation Gap: Turning a signed agreement into a structured invoice is a tedious manual process.",
         "Scope Creep: Clients frequently ask for \"one quick change\" via text. There is no structured way "
         "to bill for these small changes without a massive administrative headache.",
     ]),
    ("Admin Overhead",
     [
         "The Time Sink: Time tracking, summarizing deliverables, and reconciling payments pull creators "
         "away from billable work.",
         "The Anchor Gap: Current reviews suggest Anchor feels built for accounting firms. Solo creatives "
         "find the desktop dashboard \"heavy\" and desire more streamlined, mobile-first interactions.",
     ]),
]

for i, (title, bullets) in enumerate(pains, 1):
    section_heading(doc, f"{i}. {title}", level=2)
    for b in bullets:
        colon_idx = b.find(':')
        if colon_idx != -1:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent  = Cm(1.2)
            p.paragraph_format.first_line_indent = Cm(-0.45)
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
            rb = p.add_run("•  ")
            set_font(rb, 10.5, bold=True, color=PURPLE)
            bold_inline(p, b[:colon_idx+1], b[colon_idx+1:])
        else:
            bullet(doc, b)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 3 — MARKET LANDSCAPE
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Market Landscape")
body(doc, "To understand the opportunity, we must look at how the market currently addresses these pains.")

competitors = [
    ("1. The \"Dashboard Giants\" (HoneyBook, Bonsai, Moxie)",
     "They offer \"All-in-One\" platforms with deep automation — beautiful templates and auto-reminders.",
     "Complexity & \"Dashboard Fatigue.\" Using these tools feels like a second job. Built for the office, "
     "not for the on-the-go reality of a solo creator."),
    ("2. The \"Payment Processors\" (Stripe, FreshBooks)",
     "They focus on the plumbing: instant payouts and professional-grade invoicing.",
     "Transactional, not Relational. They are the cash register at the end of the race. They don't help "
     "negotiate or handle the awkwardness of the deal."),
    ("3. The \"AI Newcomers\" (Bookipi, Manus AI)",
     "They use AI to generate invoices from simple prompts (e.g., \"Bill $500 for a logo\").",
     "PDF Generators vs. Billing Engines. They help you write the bill, but lack the muscle to track "
     "money, enforce milestones, or manage the legal agreement lifecycle."),
]

for title, how, weakness in competitors:
    section_heading(doc, title, level=2)
    p = doc.add_paragraph()
    bold_inline(p, "How they solve it: ", how)
    para_spacing(p, before=0, after=4)
    p.paragraph_format.left_indent = Cm(0.4)
    p2 = doc.add_paragraph()
    bold_inline(p2, "The Weakness: ", weakness)
    para_spacing(p2, before=0, after=6)
    p2.paragraph_format.left_indent = Cm(0.4)

# Competitive table
section_heading(doc, "Competitive Comparison", level=3)
sp = doc.add_paragraph()
para_spacing(sp, before=0, after=4)

comp_headers = ["Feature", "The Giants (HoneyBook)", "The Processors (Stripe)", "AI Newcomers (Bookipi)"]
comp_rows = [
    ["Primary Interface",  "Desktop Dashboard",  "Mobile App / Web",  "Simple Web UI"],
    ["Agreement Speed",    "Slow (Templates)",   "Manual Entry",       "Fast (AI Prompts)"],
    ["Scope Changes",      "Manual Re-edit",     "New Invoice",        "Basic"],
    [""Bad Cop" Power",   "Medium (Reminders)", "Low (Static)",       "Low (PDF only)"],
]

tbl = doc.add_table(rows=1 + len(comp_rows), cols=4)
tbl.style = 'Table Grid'

# Header row
for i, h in enumerate(comp_headers):
    cell = tbl.cell(0, i)
    shade_cell(cell, '643CFF')
    cell_margins(cell, top=100, bottom=100, left=120, right=120)
    p = cell.paragraphs[0]
    r = p.add_run(h)
    set_font(r, 9.5, bold=True, color=WHITE)

# Data rows
for ri, row_data in enumerate(comp_rows):
    for ci, val in enumerate(row_data):
        cell = tbl.cell(ri+1, ci)
        shade_cell(cell, 'F7F5FF' if ri % 2 == 0 else 'FFFFFF')
        cell_margins(cell, top=80, bottom=80, left=120, right=120)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        set_font(r, 9.5, bold=(ci==0), color=NEAR_BLACK if ci==0 else GRAY_DARK)

sp2 = doc.add_paragraph()
para_spacing(sp2, before=4, after=4)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 4 — THE OPPORTUNITY
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, "The Opportunity")
body(doc,
    "By identifying the gaps in existing solutions — specifically Admin Overhead and Negotiation Friction "
    "— the focus lands on the SMS Copilot. It won't replace the Anchor Dashboard; it supercharges it. "
    "It acts as a remote control for the platform's most frequent actions.")

opportunities = [
    ("1. Objective Enforcement (Removing Emotional Friction)",
     "The Copilot acts as the \"Bad Cop.\" Users text their terms (e.g., \"Draft a $4k agreement, 50% "
     "upfront\"), and the AI immediately translates this into Anchor's strict, automated milestone "
     "structure. The system — not the freelancer — enforces the paywall, preserving the human relationship."),
    ("2. Real-Time Agreement Drift Detection",
     "When a client requests extra work verbally or via text, the user doesn't need to log in anywhere. "
     "They text the bot: \"Add a $300 line item for extra revisions.\" The Copilot updates the live "
     "agreement and sends an automated SMS to the client for approval. Scope creep is captured in seconds."),
    ("3. Auto-Invoicing from Notes",
     "The SMS interface removes the desktop UI entirely. At the end of the day, users can do a raw text "
     "dump: \"Did 4 hours of consulting today.\" The AI parses the text, categorises it, and queues the "
     "invoice. Admin time drops from minutes to a single text message."),
    ("4. Solving the \"Anchor Gap\" (Standardisation over Customisation)",
     "User feedback suggests Anchor can feel \"heavy\" (built for accountants). The Copilot reframes this: "
     "Standardisation is a feature. By shifting to native SMS, we provide a hyper-streamlined experience. "
     "Users don't worry about formatting because the AI guarantees a clean, branded mobile checkout every time."),
]

for title, desc in opportunities:
    section_heading(doc, title, level=2)
    body(doc, desc)

callout_box(doc,
    "Strategic Verdict:",
    "The market is currently building better tools, but freelancers want better partners. While "
    "competitors focus on \"customisation,\" Anchor wins by focusing on time-to-initiation. The SMS Copilot "
    "ensures that the moment a deal is struck, the billing engine is already running.")

page_break(doc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 5 — PRIORITISATION
# ══════════════════════════════════════════════════════════════════════════════
section_heading(doc, "Prioritisation (RICE Framework)")
body(doc,
    "To ensure the Anchor SMS Copilot addresses the most critical user needs while remaining technically "
    "feasible, the RICE Framework (Reach, Impact, Confidence, Effort) was applied. This objective scoring "
    "system allows movement past gut feelings and towards features offering the highest ROI.")

rice_headers = ["Feature", "Reach (1–10)", "Impact (1–3)", "Confidence %", "Effort (1–5)", "RICE Score"]
rice_rows = [
    ["Objective Enforcement",       "10", "3 (Massive)", "90%", "2", "15.0"],
    ["Standardisation (Solo UI)",   "9",  "2 (High)",    "80%", "1", "18.0"],
    ["Scope Drift Detection",       "7",  "3 (Massive)", "50%", "4", "4.2"],
    ["Auto-Invoicing / Notes",      "8",  "1 (Medium)",  "50%", "2", "2.8"],
]

tbl2 = doc.add_table(rows=1 + len(rice_rows), cols=6)
tbl2.style = 'Table Grid'

for i, h in enumerate(rice_headers):
    cell = tbl2.cell(0, i)
    shade_cell(cell, '3A1ECC')
    cell_margins(cell, top=90, bottom=90, left=100, right=100)
    p = cell.paragraphs[0]
    r = p.add_run(h)
    set_font(r, 9, bold=True, color=WHITE)

score_colors = {'18.0': '22C55E', '15.0': '643CFF', '4.2': 'FFA500', '2.8': 'FFA500'}
for ri, row_data in enumerate(rice_rows):
    for ci, val in enumerate(row_data):
        cell = tbl2.cell(ri+1, ci)
        shade_cell(cell, 'F4F2FA' if ri % 2 == 0 else 'FFFFFF')
        cell_margins(cell, top=80, bottom=80, left=100, right=100)
        p = cell.paragraphs[0]
        is_score = ci == 5
        r = p.add_run(val)
        if is_score and val in score_colors:
            hex_c = score_colors[val]
            set_font(r, 9.5, bold=True,
                     color=RGBColor(int(hex_c[0:2],16), int(hex_c[2:4],16), int(hex_c[4:6],16)))
        else:
            set_font(r, 9.5, bold=(ci==0), color=NEAR_BLACK if ci==0 else GRAY_DARK)

sp3 = doc.add_paragraph()
para_spacing(sp3, before=4, after=2)

rationale = [
    ("P0 — SMS-to-Checkout Standardisation (Score 18.0)",
     "Highest score due to Reach and Low Effort. Solves Anchor's brand perception gap for 100% of users "
     "immediately. Engineering effort is minimal because we're simplifying output, not building complex "
     "customisation tools. Immediate priority for MVP."),
    ("P1 — Objective Enforcement (Score 15.0)",
     "The \"heart\" of the SMS Copilot with the highest Impact score (3). Turns Anchor from a tool into a "
     "partner. 90% confidence backed by the #1 user pain point: late payments and the \"begging\" factor."),
    ("Phase 2 — Scope Drift & Auto-Invoicing",
     "High theoretical value but sit lower due to Technical Confidence and Effort. Scope Drift is complex "
     "(Effort 4). Auto-Invoicing requires the highest degree of user behaviour change, with only 50% "
     "confidence in proactive usage."),
]

for title, desc in rationale:
    section_heading(doc, title, level=2)
    body(doc, desc)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 6 — THE SOLUTION (PROTOTYPE)
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_heading(doc, "The Solution: Prototype Walkthrough")
body(doc,
    "The prototype is a functional React application that demonstrates the full SMS Copilot loop. It "
    "renders two phone screens side by side — the freelancer's phone (the Copilot interface) and the "
    "client's phone (the receiving end) — alongside a live mirror of the Anchor dashboard. Every text "
    "message the freelancer sends triggers a real platform action. The dashboard updates instantly, "
    "with no manual input required.")

section_heading(doc, "Architecture", level=2)

arch_layers = [
    ("1. The SMS Interface",
     "A mobile-first chat UI where the freelancer types natural language commands. Messages are sent to "
     "a backend API powered by Claude AI, which interprets intent and returns both an SMS reply and a "
     "structured action payload."),
    ("2. The Action Engine",
     "A stateful context layer that receives the AI's action payload and mutates platform data "
     "accordingly. It supports six action types: draft_agreement, confirm, add_line_item, mark_paid, "
     "send_reminder, and update_agreement."),
    ("3. The Platform Mirror",
     "A real-time replica of Anchor's Agreements dashboard. Every SMS action is immediately reflected "
     "in the table — new agreements appear, statuses change, milestones update — demonstrating that the "
     "Copilot is not a standalone tool but a remote control for the existing platform."),
]

for title, desc in arch_layers:
    p = doc.add_paragraph()
    para_spacing(p, before=6, after=2)
    p.paragraph_format.left_indent = Cm(0.5)
    rb = p.add_run(f"{title}  ")
    set_font(rb, 10.5, bold=True, color=PURPLE)
    r2 = p.add_run(desc)
    set_font(r2, 10.5, color=GRAY_DARK)

section_heading(doc, "Core User Flows", level=2)

flows = [
    (
        'Flow 1: Draft a New Agreement',
        '"Draft a $4k agreement with Rachel, 50% upfront"',
        [
            "The freelancer texts the Copilot immediately after verbally agreeing on project scope.",
            "The AI parses the total, client name, and payment split, then replies: \\"Ready to send: "
            "Rachel, $4,000 — $2,000 upfront + $2,000 on delivery. Reply YES to confirm.\\"",
            "The freelancer replies YES.",
            "The Copilot creates a live agreement in Anchor's system and fires an automated SMS to the "
            "client: \\"Hi Rachel! Anchor sent you a payment agreement for $4,000. Tap to review and approve.\\"",
            "The client taps YES. Milestone 1 flips to \"paid.\" Freelancer receives: \\"Rachel paid $2,000. "
            "Milestone 2 unlocked.\\"",
        ],
        "Time from deal to billing engine running: under 60 seconds.",
    ),
    (
        'Flow 2: Capture Scope Creep',
        '"Add $300 for extra revisions"',
        [
            "A client asks for an extra round of edits over text or call.",
            "The freelancer texts the Copilot the change — no dashboard required.",
            "The AI adds a $300 — Extra revisions line item to the active agreement and notifies the "
            "client: \\"New line item: 'Extra revisions' — $300. Reply YES to approve.\\"",
            "The client replies YES. The line item is approved and the agreement total is updated.",
        ],
        "Scope creep is captured in seconds, not days. No awkward conversation required.",
    ),
    (
        'Flow 3: Send Payment Reminders',
        '"Send David a reminder" or "Remind all overdue clients"',
        [
            "The freelancer texts the Copilot from anywhere — after a meeting, on a commute.",
            "The Copilot identifies the overdue agreement and fires an automated reminder to the client.",
            "The dashboard logs a \"last reminded\" timestamp, creating an audit trail.",
        ],
        "The freelancer never has to write a \"just checking in\" message again. The system is the bad cop.",
    ),
    (
        'Flow 4: Check Portfolio Status',
        '"Who owes me money?" / "How much have I collected?"',
        [
            "The freelancer queries their full portfolio state at any time with plain English.",
            "The Copilot aggregates across all active agreements and replies: \\"David Lee owes $7,200 "
            "(overdue since Feb 28) and Maya Cohen owes $1,900 (pending approval). Total outstanding: $9,100.\\"",
        ],
        "A full financial snapshot, delivered as a single text message.",
    ),
]

for flow_title, prompt, steps, takeaway in flows:
    section_heading(doc, flow_title, level=3)

    # Prompt callout
    tbl_q = doc.add_table(rows=1, cols=1)
    no_borders(tbl_q)
    qcell = tbl_q.cell(0, 0)
    shade_cell(qcell, 'F0EBFF')
    cell_margins(qcell, top=80, bottom=80, left=140, right=140)
    set_cell_border(qcell,
        left={'val':'single','sz':16,'color':'643CFF'},
        top={'val':'nil'}, bottom={'val':'nil'}, right={'val':'nil'})
    qp = qcell.paragraphs[0]
    rq = qp.add_run(prompt)
    set_font(rq, 10.5, italic=True, color=PURPLE_DARK)
    para_spacing(qp, before=0, after=0)
    doc.add_paragraph()

    for i, step in enumerate(steps, 1):
        p = doc.add_paragraph()
        p.paragraph_format.left_indent      = Cm(1.2)
        p.paragraph_format.first_line_indent = Cm(-0.5)
        p.paragraph_format.space_before     = Pt(2)
        p.paragraph_format.space_after      = Pt(2)
        rn = p.add_run(f"{i}.  ")
        set_font(rn, 10.5, bold=True, color=PURPLE)
        rs = p.add_run(step)
        set_font(rs, 10.5, color=GRAY_DARK)

    sp_f = doc.add_paragraph()
    para_spacing(sp_f, before=4, after=2)
    callout_box(doc, "Key insight:", takeaway, bg='EDFFF4', accent='22C55E')

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 7 — RISKS
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_heading(doc, "Risks")

risks = [
    (
        "1. AI Reliability — Misinterpretation of Intent",
        "Medium",
        "FFA500",
        "The Copilot depends on a language model to correctly parse financial instructions. A misread — "
        "turning \"change Maya's total to $2,200\" into a different action — could corrupt a live agreement.",
        [
            "Every destructive or financial action requires a confirmation step before execution.",
            "The AI is instructed to show a preview and ask for approval before writing to the system.",
            "Ambiguous inputs default to clarification, not assumption.",
        ],
    ),
    (
        "2. SMS as a Channel — Delivery & Trust",
        "Medium",
        "FFA500",
        "Standard SMS has no delivery guarantees. A reminder that silently fails gives the freelancer "
        "false confidence. Receiving a payment request from an unknown number creates phishing-like "
        "friction for the client.",
        [
            "Near-term: Use a verified Twilio short code or branded sender ID. Include the freelancer's "
            "name in every outbound message.",
            "Long-term: Offer WhatsApp or iMessage Business as higher-trust delivery channels.",
        ],
    ),
    (
        "3. Adoption Friction — Changing User Behaviour",
        "High",
        "FF4D4D",
        "The biggest risk is not technical — it is behavioural. Asking freelancers to manage billing "
        "via text when they already have a dashboard requires a strong habit loop. Users who onboard but "
        "never form the habit will churn quickly.",
        [
            "The Copilot's value must be felt in the first two minutes.",
            "Onboarding should immediately demonstrate the \"deal struck → agreement live\" scenario.",
            "The prototype's demo script — three messages from draft to scope update — is the "
            "designed onboarding moment.",
            "Key early signals: message-to-action conversion rate and time-to-first-send.",
        ],
    ),
    (
        "4. Data Security & Consent",
        "Low – Medium",
        "22C55E",
        "Billing data, client names, and payment amounts flowing through SMS and an AI model creates "
        "regulatory exposure under GDPR and CCPA.",
        [
            "The AI model processes only data shared in each session — no persistent storage of raw "
            "message content.",
            "Client phone numbers are stored within Anchor's existing infrastructure.",
            "Clear consent language is required at client onboarding.",
            "Marginal risk is the AI processing layer, which requires a DPA with the model provider.",
        ],
    ),
    (
        "5. Competitive Response",
        "Low at launch",
        "22C55E",
        "If the feature gains traction, HoneyBook or Bonsai could build a comparable SMS layer within "
        "6–12 months. The moat is not the SMS interface — it is the depth of Anchor's billing engine behind it.",
        [
            "A competitor building an SMS wrapper on a weaker billing engine gets a chatbot, not a copilot.",
            "Anchor's milestone enforcement, legal agreement infrastructure, and payment rails are the "
            "real differentiator.",
        ],
    ),
]

for title, risk_level, risk_hex, context_text, mitigations in risks:
    section_heading(doc, title, level=2)

    # Risk level badge
    tbl_r = doc.add_table(rows=1, cols=2)
    no_borders(tbl_r)
    tbl_r.columns[0].width = Cm(3)

    badge_cell = tbl_r.cell(0, 0)
    shade_cell(badge_cell, risk_hex)
    cell_margins(badge_cell, top=60, bottom=60, left=100, right=100)
    bp = badge_cell.paragraphs[0]
    bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    br = bp.add_run(f"Risk: {risk_level}")
    set_font(br, 9, bold=True, color=WHITE)
    para_spacing(bp, before=0, after=0)

    ctx_cell = tbl_r.cell(0, 1)
    cell_margins(ctx_cell, top=60, bottom=60, left=140, right=60)
    cp2 = ctx_cell.paragraphs[0]
    rc2 = cp2.add_run(context_text)
    set_font(rc2, 10, italic=True, color=GRAY_MID)
    para_spacing(cp2, before=0, after=0)

    sp_r = doc.add_paragraph()
    para_spacing(sp_r, before=4, after=2)

    mit_label = doc.add_paragraph()
    rm = mit_label.add_run("Mitigations:")
    set_font(rm, 10.5, bold=True, color=NEAR_BLACK)
    para_spacing(mit_label, before=2, after=2)

    for m in mitigations:
        bullet(doc, m)

    sp_end = doc.add_paragraph()
    para_spacing(sp_end, before=2, after=2)

# ══════════════════════════════════════════════════════════════════════════════
# SECTION 8 — BUSINESS MODEL
# ══════════════════════════════════════════════════════════════════════════════
page_break(doc)
section_heading(doc, "Business Model")

section_heading(doc, "Positioning: The Copilot as a Premium Add-On", level=2)
body(doc,
    "The SMS Copilot should not be bundled into Anchor's base plan at launch. It is a distinct, "
    "high-value capability with clear willingness-to-pay: freelancers who use it will recover its cost "
    "on the first single reminder that gets a late invoice paid.")
body(doc,
    "The model is an optional subscription layer on top of a user's existing Anchor account.")

section_heading(doc, "Pricing Tiers", level=2)

tier_headers = ["Tier", "Name", "Price", "Who It's For", "What's Included"]
tier_rows = [
    ["Free",     "Anchor Core",       "$0 / mo",   "New & light users",
     "Full dashboard access. No SMS Copilot."],
    ["Starter",  "Copilot Starter",   "$12 / mo",  "Freelancers, 1–5 active clients",
     "50 Copilot messages/mo. Agreement drafting, reminders."],
    ["Pro",      "Copilot Pro",       "$29 / mo",  "Active freelancers & small agencies",
     "Unlimited messages. Scope drift, line items, portfolio snapshots."],
    ["Business", "Copilot Business",  "$59 / mo",  "Agencies, 2–5 team members",
     "Everything in Pro + multi-user, team agreements, priority AI routing."],
]

tier_bg = ['FFFFFF', 'F0EBFF', 'EDE4FF', 'DDD4FF']

tbl3 = doc.add_table(rows=1 + len(tier_rows), cols=5)
tbl3.style = 'Table Grid'

for i, h in enumerate(tier_headers):
    cell = tbl3.cell(0, i)
    shade_cell(cell, '3A1ECC')
    cell_margins(cell, top=90, bottom=90, left=110, right=110)
    p = cell.paragraphs[0]
    r = p.add_run(h)
    set_font(r, 9.5, bold=True, color=WHITE)

for ri, row_data in enumerate(tier_rows):
    for ci, val in enumerate(row_data):
        cell = tbl3.cell(ri+1, ci)
        shade_cell(cell, tier_bg[ri])
        cell_margins(cell, top=90, bottom=90, left=110, right=110)
        p = cell.paragraphs[0]
        r = p.add_run(val)
        is_price = ci == 2
        set_font(r, 9.5,
                 bold=(ci in (0, 1, 2)),
                 color=PURPLE_DARK if is_price else (NEAR_BLACK if ci < 3 else GRAY_DARK))

sp4 = doc.add_paragraph()
para_spacing(sp4, before=6, after=2)

section_heading(doc, "Revenue Logic", level=2)

rev_points = [
    ("Why $12–$29 is the right range",
     "The average freelancer invoice is $3,000–$8,000. One recovered late payment — a $3,600 overdue "
     "balance — is worth 25 years of the Starter plan. Users will pay to not feel the emotional cost "
     "of chasing money. This is a stress-reduction product, not a \"nice to have.\""),
    ("Message Volume as a Usage Governor",
     "The Starter plan's 50-message cap is generous for light users (~10–15 client interactions per "
     "month) but creates a natural upgrade path for anyone running more than 5 active agreements. "
     "Pro should feel obvious by month two, not restrictive."),
    ("Potential Revenue Ceiling",
     "If the Copilot converts 20% of Anchor's active user base to a paid tier at $12/mo Starter, "
     "it represents meaningful ARR expansion on top of existing transaction revenue. At $29/mo Pro, "
     "a 5,000-user cohort generates $1.74M ARR."),
]

for title, desc in rev_points:
    section_heading(doc, title, level=3)
    body(doc, desc)

section_heading(doc, "Monetisation Signals to Watch (First 90 Days)", level=2)

metrics = [
    ("Time-to-first-send",
     "If a new user sends their first Copilot message within the first session, the product is working."),
    ("Reminder-to-payment conversion",
     "What percentage of AI-sent reminders result in a payment within 48 hours? This is the headline "
     "ROI metric for user retention and word-of-mouth."),
    ("Scope drift capture rate",
     "How often is the \"add a line item\" flow used per active agreement? High usage signals the feature "
     "is genuinely embedded in workflow."),
    ("Starter → Pro upgrade rate",
     "If fewer than 15% of Starter users upgrade to Pro within 60 days, the message cap may be too "
     "high or Pro features are not compelling enough."),
]

tbl4 = doc.add_table(rows=1 + len(metrics), cols=2)
no_borders(tbl4)

# header
for i, h in enumerate(["Metric", "What It Tells Us"]):
    cell = tbl4.cell(0, i)
    shade_cell(cell, 'F0EBFF')
    cell_margins(cell, top=80, bottom=80, left=120, right=120)
    thin_bottom(cell, 'C4BAFF')
    p = cell.paragraphs[0]
    r = p.add_run(h)
    set_font(r, 10, bold=True, color=PURPLE_DARK)

for ri, (m_title, m_desc) in enumerate(metrics):
    for ci, val in enumerate([m_title, m_desc]):
        cell = tbl4.cell(ri+1, ci)
        shade_cell(cell, 'FAFAFA' if ri % 2 == 0 else 'FFFFFF')
        cell_margins(cell, top=80, bottom=80, left=120, right=120)
        thin_bottom(cell, 'EEEBF8')
        p = cell.paragraphs[0]
        r = p.add_run(val)
        set_font(r, 10, bold=(ci==0), color=NEAR_BLACK if ci==0 else GRAY_DARK)

sp5 = doc.add_paragraph()
para_spacing(sp5, before=6, after=4)

callout_box(doc,
    "Strategic Note:",
    "Beyond direct revenue, the Copilot dramatically raises Anchor's switching cost. A freelancer whose "
    "entire billing history, agreement lifecycle, and client communication trail lives inside Anchor's "
    "SMS thread has a compelling reason to stay. Every message sent through the Copilot is a data point "
    "that deepens the platform's understanding of the user's client relationships — a capability no "
    "competitor can replicate by copying the interface alone.")

# ── Save ───────────────────────────────────────────────────────────────────────
out_path = r"C:\Users\yamco\Desktop\Projects\Product Yam\Portfolio\Submissions\Anchor\Anchor SMS Copilot.docx"
doc.save(out_path)
print(f"Saved: {out_path}")
